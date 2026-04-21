# document_service/extractors.py

import csv
from pathlib import Path

from document_service.clients.ocr_client import run_remote_ocr
from document_service.identity import (
    build_content_hash,
    build_portable_document_id,
    build_source_ref,
)
from document_service.schemas import ExtractedDocument


def is_image_file(file_path: str) -> bool:
    """
    Check whether a path is a supported image type.
    """
    suffix = Path(file_path).suffix.lower()
    return suffix in {".png", ".jpg", ".jpeg", ".bmp", ".webp", ".tiff"}


def _read_plain_text_file(file_path: str) -> str:
    """
    Read common text-based files.
    """
    return Path(file_path).read_text(encoding="utf-8", errors="ignore").strip()


def _read_csv_as_text(file_path: str) -> str:
    """
    Convert CSV into a readable plain-text representation.
    """
    lines = []

    with open(file_path, "r", encoding="utf-8", errors="ignore", newline="") as f:
        reader = csv.reader(f)
        for row in reader:
            lines.append(" | ".join(cell.strip() for cell in row))

    return "\n".join(lines).strip()


def _read_docx_text(file_path: str) -> str:
    """
    Read DOCX using python-docx if installed.
    """
    try:
        from docx import Document
    except Exception as e:
        raise RuntimeError(
            "python-docx is not installed. Install it to read DOCX files."
        ) from e

    doc = Document(file_path)
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    return "\n".join(paragraphs).strip()


def extract_non_pdf_document(file_path: str, enable_ocr: bool = True) -> ExtractedDocument:
    """
    Extract non-PDF documents.

    Handles:
    - txt
    - md
    - py
    - json
    - yaml
    - yml
    - log
    - csv
    - docx
    - image files (via OCR service)

    Phase 4 change:
    - document identity is built from extracted content + file naming context,
      not from the absolute local path
    """
    path = Path(file_path)

    if not path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")

    if not path.is_file():
        raise ValueError(f"Path is not a file: {file_path}")

    extension = path.suffix.lower()
    file_name = path.name

    metadata = {
        "source_path": str(path.resolve()),
        "size_bytes": path.stat().st_size,
    }

    text = ""
    ocr_used = False

    if extension in {".txt", ".md", ".py", ".json", ".yaml", ".yml", ".log"}:
        text = _read_plain_text_file(file_path)

    elif extension == ".csv":
        text = _read_csv_as_text(file_path)

    elif extension == ".docx":
        text = _read_docx_text(file_path)

    elif is_image_file(file_path):
        if not enable_ocr:
            raise ValueError("OCR is disabled, so image files cannot be processed.")

        ocr_result = run_remote_ocr(file_path)
        if not ocr_result.get("success"):
            raise RuntimeError(ocr_result.get("message", "OCR failed."))

        text = (ocr_result.get("text") or "").strip()
        ocr_used = True
        metadata["ocr_mode"] = "image_file"

    else:
        raise ValueError(
            f"Unsupported non-PDF file type: {extension}. "
            "Supported: txt, md, py, json, yaml, yml, log, csv, docx, image files."
        )

    text = text.strip()
    content_hash = build_content_hash(text)
    document_id = build_portable_document_id(
        file_name=file_name,
        extension=extension,
        text=text,
    )

    return ExtractedDocument(
        document_id=document_id,
        source_ref=build_source_ref(file_path),
        content_hash=content_hash,
        local_file_path=str(path.resolve()),
        file_name=file_name,
        extension=extension,
        text=text,
        metadata=metadata,
        ocr_used=ocr_used,
    )