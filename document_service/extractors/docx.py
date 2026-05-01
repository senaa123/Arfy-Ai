"""
DOCX extraction.

Kept separate so DOCX support remains explicit and easy to maintain.
"""

from __future__ import annotations

from pathlib import Path

from document_service.identity import (
    build_content_hash,
    build_portable_document_id,
    build_source_ref,
)
from document_service.schemas import ExtractedDocument


def _read_docx_text(file_path: str | Path) -> str:
    """
    Read DOCX using python-docx if available.
    """
    try:
        from docx import Document
    except Exception as e:
        raise RuntimeError(
            "python-docx is not installed. Install it to read DOCX files."
        ) from e

    doc = Document(file_path)
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    return "\n".join(paragraphs).strip()


def extract_docx_document(file_path: str, enable_ocr: bool = True) -> ExtractedDocument:
    """
    Extract DOCX into the normalized ExtractedDocument shape.

    OCR is not used here, but the parameter stays for a consistent interface.
    """
    _ = enable_ocr

    path = Path(file_path)

    if not path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")

    if not path.is_file():
        raise ValueError(f"Path is not a file: {file_path}")

    if path.suffix.lower() != ".docx":
        raise ValueError(f"Expected .docx file, got: {path.suffix.lower()}")

    text = _read_docx_text(path)
    content_hash = build_content_hash(text)
    document_id = build_portable_document_id(
        file_name=path.name,
        extension=".docx",
        text=text,
    )

    metadata = {
        "source_path": str(path.resolve()),
        "size_bytes": path.stat().st_size,
    }

    return ExtractedDocument(
        document_id=document_id,
        source_ref=build_source_ref(file_path),
        content_hash=content_hash,
        local_file_path=str(path.resolve()),
        file_name=path.name,
        extension=".docx",
        text=text,
        metadata=metadata,
        ocr_used=False,
    )
